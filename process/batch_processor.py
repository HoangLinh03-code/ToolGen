import time
import re
from typing import List, Dict, Tuple
from dataclasses import dataclass
from api.callAPI import VertexClient
from process.ques_valid import QuestionValidator, ValidQuestionStorage
from process.PDF_Scan import enhance_prompt_with_pdf_scan
from docx import Document
from concurrent.futures import ThreadPoolExecutor, as_completed

@dataclass
class OptimizedBatchConfig:
    """Cấu hình tối ưu cho tốc độ"""
    batch_size: int = 20  # TĂNG từ 10 → 20 câu (giảm 50% số lần gọi API)
    max_retry_per_batch: int = 3  # GIẢM từ 10 → 3 (chỉ retry khi thực sự cần)
    max_write_retry: int = 2  # GIẢM từ 3 → 2
    max_final_recovery_retry: int = 5  # GIẢM từ 10 → 5
    parallel_image_generation: bool = True  # BẬT song song hóa
    skip_minor_validation: bool = True  # BỎ QUA validation không quan trọng


class BatchProcessorOptimized:
    """
    Processor tối ưu cho tốc độ
    Mục tiêu: Nhanh gấp 5-10 lần version cũ
    """
    
    def __init__(self, project_id: str, creds, model_name: str = "gemini-2.5-pro"):
        self.client = VertexClient(project_id, creds, model_name)
        self.project_id = project_id
        self.creds = creds
        self.config = OptimizedBatchConfig()
        self.image_tracker = None
        
    def process_all_batches_fast(
        self,
        prompt_content: str,
        pdf_files: List[str],
        question_type: str = "tracnghiem",
        doc: Document = None
    ) -> Tuple[ValidQuestionStorage, List[int]]:
        required_count = 80 if question_type == "tracnghiem" else 40
        storage = ValidQuestionStorage()
        all_failed_nums = []
        
        print(f"\n{'='*70}")
        print(f"🚀 BẮT ĐẦU XỬ LÝ NHANH - {required_count} CÂU")
        print(f"⚡ Batch size: {self.config.batch_size} (lớn hơn 2x)")
        print(f"⚡ Max retry: {self.config.max_retry_per_batch} (ít hơn 3x)")
        print(f"{'='*70}\n")
        
        # ============ SCAN PDF (GIỮ NGUYÊN) ============
        print("📖 Bước 1: Scan PDF...")
        start_scan = time.time()
        enhanced_prompt = enhance_prompt_with_pdf_scan(
            prompt_content, pdf_files, self.project_id, self.creds
        )
        print(f"✅ Scan xong ({time.time() - start_scan:.1f}s)\n")
        
        # ============ KHỞI TẠO IMAGE TRACKER ============
        from process.text2Image import calculate_optimal_image_count, ImageGenerationTracker
        target_image_count = calculate_optimal_image_count(required_count, 0.17)
        self.image_tracker = ImageGenerationTracker(target_image_count)
        
        # ============ TÍNH SỐ BATCH TỐI ƯU ============
        num_batches = (required_count + self.config.batch_size - 1) // self.config.batch_size
        print(f"📦 Tổng batch: {num_batches} (mỗi batch {self.config.batch_size} câu)\n")
        
        # ============ PHASE 1: XỬ LÝ TỪNG BATCH - TỐI ƯU ============
        print(f"\n{'='*70}")
        print("🔥 PHASE 1: XỬ LÝ CÁC BATCH (OPTIMIZED)")
        print(f"{'='*70}\n")
        
        total_start_time = time.time()
        
        for batch_idx in range(num_batches):
            batch_start_time = time.time()
            
            batch_start = batch_idx * self.config.batch_size + 1
            batch_end = min((batch_idx + 1) * self.config.batch_size, required_count)
            batch_nums = list(range(batch_start, batch_end + 1))
            
            print(f"\n{'='*70}")
            print(f"📦 BATCH {batch_idx + 1}/{num_batches}: Câu {batch_start}-{batch_end}")
            print(f"{'='*70}\n")
            
            # XỬ LÝ BATCH TỐI ƯU
            failed_in_batch = self._process_single_batch_optimized(
                batch_nums,
                enhanced_prompt,
                storage,
                question_type,
                doc
            )
            
            all_failed_nums.extend(failed_in_batch)
            
            batch_time = time.time() - batch_start_time
            print(f"\n✅ Hoàn thành Batch {batch_idx + 1} ({batch_time:.1f}s)")
            print(f"   - Thành công: {len(batch_nums) - len(failed_in_batch)}/{len(batch_nums)}")
            if failed_in_batch:
                print(f"   - Thất bại: {failed_in_batch}")
        
        total_time = time.time() - total_start_time
        print(f"\n⏱️  Tổng thời gian Phase 1: {total_time:.1f}s")
        print(f"⚡ Tốc độ: {total_time/required_count:.2f}s/câu")
        
        # ============ PHASE 2: RECOVERY (CHỈ KHI THẬT SỰ CẦN) ============
        if all_failed_nums and len(all_failed_nums) < required_count * 0.2:  # < 20% thất bại
            print(f"\n{'='*70}")
            print("🔄 PHASE 2: RECOVERY - XỬ LÝ CÂU THẤT BẠI")
            print(f"{'='*70}\n")
            
            print(f"⚠️  Có {len(all_failed_nums)} câu thất bại")
            print(f"📋 Danh sách: {all_failed_nums[:10]}")
            
            recovery_start = time.time()
            final_failed = self._final_recovery_optimized(
                all_failed_nums,
                enhanced_prompt,
                storage,
                question_type,
                doc
            )
            recovery_time = time.time() - recovery_start
            print(f"\n⏱️  Recovery time: {recovery_time:.1f}s")
            
            all_failed_nums = final_failed
        elif len(all_failed_nums) >= required_count * 0.2:
            print(f"\n⚠️  Quá nhiều câu thất bại ({len(all_failed_nums)}/{required_count})")
            print(f"💡 Gợi ý: Kiểm tra prompt hoặc PDF")
        
        # ============ TỔNG KẾT ============
        total_elapsed = time.time() - total_start_time
        success_count = storage.get_valid_count()
        
        print(f"\n{'='*70}")
        print(f"📊 TỔNG KẾT")
        print(f"{'='*70}")
        print(f"✅ Câu đã ghi: {success_count}/{required_count}")
        print(f"❌ Thất bại cuối cùng: {len(all_failed_nums)}")
        print(f"⏱️  Tổng thời gian: {total_elapsed:.1f}s")
        print(f"⚡ Tốc độ trung bình: {total_elapsed/required_count:.2f}s/câu")
        
        if all_failed_nums:
            print(f"\n📋 Danh sách thất bại: {all_failed_nums}")
        else:
            print(f"\n🎉 100% THÀNH CÔNG!")
        
        if self.image_tracker:
            print(f"\n🖼️  THỐNG KÊ ẢNH:")
            print(f"   - Đã sinh: {self.image_tracker.generated_count}/{target_image_count}")
            print(f"   - Placeholder: {self.image_tracker.placeholder_count}")
        
        print(f"{'='*70}\n")
        
        return storage, all_failed_nums
    
    def _process_single_batch_optimized(
        self,
        batch_nums: List[int],
        enhanced_prompt: str,
        storage: ValidQuestionStorage,
        question_type: str,
        doc: Document
    ) -> List[int]:
        """
        XỬ LÝ BATCH TỐI ƯU
        
        Cải tiến:
        1. Giảm retry từ 10 → 3
        2. Validation nhẹ hơn
        3. Sinh + Ghi liền, không qua storage
        """
        print(f"\n{'─'*70}")
        print(f"🔥 BẮT ĐẦU XỬ LÝ BATCH OPTIMIZED")
        print(f"📋 Câu cần xử lý: {batch_nums}")
        print(f"{'─'*70}\n")
        
        # ============ PHASE 1: SINH NHANH ============
        print(f"🚀 PHASE 1: SINH CÂU HỎI")
        print(f"{'─'*60}\n")
        
        max_attempts = self.config.max_retry_per_batch
        
        for attempt in range(1, max_attempts + 1):
            missing = [n for n in batch_nums if not storage.has_question(n)]
            
            if not missing:
                print(f"🎉 Đã đủ {len(batch_nums)} câu!\n")
                break
            
            print(f"\n🔄 Lần thử {attempt}/{max_attempts}")
            print(f"📝 Cần sinh: {len(missing)} câu: {missing}\n")
            
            # TẠO PROMPT
            batch_prompt = self._create_batch_prompt_fast(
                missing,
                enhanced_prompt,
                storage,
                question_type
            )
            
            # GỌI API
            try:
                print(f"📤 Gửi request...")
                start_time = time.time()
                
                response = self.client.send_data_to_check(
                    prompt=batch_prompt,
                    temperature=0.45 + (attempt * 0.05)
                )
                
                api_time = time.time() - start_time
                print(f"✅ Nhận response ({len(response)} ký tự, {api_time:.1f}s)\n")
                
                # VALIDATE VÀ LƯU NHANH
                newly_valid = self._validate_and_store_fast(
                    response,
                    missing,
                    storage,
                    question_type
                )
                
                print(f"\n📊 Kết quả lần {attempt}:")
                print(f"   ✅ Valid: {len(newly_valid)}")
                print(f"   ❌ Còn thiếu: {len(missing) - len(newly_valid)}")
                
            except Exception as e:
                print(f"❌ Lỗi: {str(e)}\n")
                time.sleep(1)  # Giảm delay từ 3s → 1s
        
        valid_in_storage = [n for n in batch_nums if storage.has_question(n)]
        missing_in_storage = [n for n in batch_nums if not storage.has_question(n)]
        
        print(f"\n{'─'*60}")
        print(f"📊 KẾT QUẢ PHASE 1:")
        print(f"   ✅ Có trong storage: {len(valid_in_storage)}/{len(batch_nums)}")
        print(f"   ❌ Vẫn thiếu: {len(missing_in_storage)}")
        print(f"{'─'*60}\n")
        
        # ============ PHASE 2: GHI VÀO DOCX NHANH ============
        print(f"\n✍️  PHASE 2: GHI VÀO DOCX")
        print(f"{'─'*60}\n")
        
        if not valid_in_storage:
            return batch_nums
        
        valid_in_storage.sort()
        failed_to_write = []
        
        for qnum in valid_in_storage:
            print(f"✍️  Ghi câu {qnum}...")
            
            qtext = storage.valid_questions[qnum]
            write_success = False
            
            for write_attempt in range(1, self.config.max_write_retry + 1):
                try:
                    self._write_question_to_doc_fast(
                        qnum, qtext, doc, question_type,
                        image_tracker=self.image_tracker
                    )
                    
                    print(f"   ✅ OK")
                    write_success = True
                    break
                    
                except Exception as e:
                    print(f"   ❌ Lỗi: {str(e)[:40]}")
                    
                    if write_attempt < self.config.max_write_retry:
                        # REGEN NHANH
                        new_qtext = self._regenerate_single_fast(
                            qnum, enhanced_prompt, storage, question_type
                        )
                        if new_qtext:
                            storage.replace_question(qnum, new_qtext)
                            qtext = new_qtext
                            time.sleep(0.5)  # Giảm delay
                    else:
                        break
            
            if not write_success:
                failed_to_write.append(qnum)
        
        print(f"\n{'─'*60}")
        print(f"📊 KẾT QUẢ PHASE 2:")
        print(f"   ✅ Đã ghi: {len(valid_in_storage) - len(failed_to_write)}")
        print(f"   ❌ Thất bại: {len(failed_to_write)}")
        print(f"{'─'*60}\n")
        
        total_failed = list(set(missing_in_storage + failed_to_write))
        return total_failed
    
    def _validate_and_store_fast(
        self,
        response: str,
        expected_nums: List[int],
        storage: ValidQuestionStorage,
        question_type: str
    ) -> List[int]:
        """VALIDATION NHANH - BỎ QUA CHECK KHÔNG QUAN TRỌNG"""
        validator = QuestionValidator(question_type=question_type)
        questions = validator.parse_questions(response)
        
        newly_valid = []
        
        for qnum, qtext in sorted(questions.items()):
            if qnum not in expected_nums:
                continue
            
            # VALIDATION NHẸ - CHỈ CHECK THIẾT YẾU
            if self.config.skip_minor_validation:
                # CHỈ CHECK: có title, có content, có đáp án
                has_basic = (
                    "**Câu" in qtext and
                    len(qtext) > 100 and
                    qtext.count("A.") >= 1
                )
                
                if has_basic:
                    storage.replace_question(qnum, qtext)
                    newly_valid.append(qnum)
                    print(f"   ✅ Câu {qnum}: Quick validation OK")
                else:
                    print(f"   ❌ Câu {qnum}: Thiếu phần cơ bản")
            else:
                # VALIDATION ĐẦY ĐỦ (chậm hơn)
                validation = validator.validate_single_question(qtext, qnum)
                if validation.is_valid:
                    storage.replace_question(qnum, qtext)
                    newly_valid.append(qnum)
                    print(f"   ✅ Câu {qnum}: Full validation OK")
                else:
                    print(f"   ❌ Câu {qnum}: {', '.join(validation.missing_parts)}")
        
        return newly_valid
    
    def _write_question_to_doc_fast(
        self,
        qnum: int,
        qtext: str,
        doc: Document,
        question_type: str,
        image_tracker=None
    ):
        """GHI DOCX NHANH - TỐI THIỂU XỬ LÝ"""
        from process.response2docx import (
            is_question_title, is_answer_option, is_solution_header,
            is_correct_answer, is_separator, is_image_line,
            extract_image_description, process_text, process_bold_text,
            handle_image_generation
        )
        from docx.shared import Pt
        
        lines = qtext.split("\n")
        
        # PARSE NHANH - KHÔNG VALIDATE
        title_line = None
        content_lines = []
        image_desc = None
        answers = []
        solution_header = None
        correct_answer = None
        explanation_lines = []
        
        in_content = False
        in_answers = False
        in_solution = False
        in_explanation = False
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if is_question_title(line):
                title_line = line.replace("**", "").strip()
                in_content = True
                continue
            
            if is_image_line(line):
                image_desc = extract_image_description(line)
                continue
            
            if is_answer_option(line):
                answers.append(line)
                in_content = False
                in_answers = True
                continue
            
            if is_solution_header(line):
                solution_header = line.replace("**", "").strip()
                in_answers = False
                in_solution = True
                continue
            
            if in_solution and is_correct_answer(line):
                correct_answer = line.strip()
                continue
            
            if is_separator(line):
                in_solution = False
                in_explanation = True
                continue
            
            if in_explanation and line:
                explanation_lines.append(line)
                continue
            
            if in_content and not in_answers and not in_solution:
                content_lines.append(line)
        
        # GHI NHANH - KHÔNG CHECK
        if title_line:
            p = doc.add_paragraph()
            run = p.add_run(title_line)
            run.bold = True
            run.font.size = Pt(12)
        
        for line in content_lines:
            if line.strip():
                p = doc.add_paragraph()
                if "**" in line:
                    process_bold_text(line, p)
                else:
                    process_text(line, p)
        
        # HÌNH ẢNH - SONG SONG HÓA
        if image_desc:
            should_generate = False
            if image_tracker:
                should_generate = image_tracker.should_generate(qnum, 80)
            
            try:
                handle_image_generation(
                    image_desc, doc,
                    attempt_generate=should_generate,
                    tracker=image_tracker
                )
            except:
                pass  # Bỏ qua lỗi ảnh để không ảnh hưởng tốc độ
        
        for ans in answers[:4]:
            p = doc.add_paragraph()
            process_text(ans, p)
        
        doc.add_paragraph()
        
        if solution_header:
            p = doc.add_paragraph()
            run = p.add_run(solution_header)
            run.bold = True
        
        if correct_answer:
            p = doc.add_paragraph()
            run = p.add_run(correct_answer)
            run.bold = True
        
        p = doc.add_paragraph()
        run = p.add_run("####")
        run.bold = True
        
        max_lines = 5 if question_type == "tracnghiem" else len(explanation_lines)
        for line in explanation_lines[:max_lines]:
            if line.strip():
                p = doc.add_paragraph()
                if "**" in line:
                    process_bold_text(line, p)
                else:
                    process_text(line, p)
    
    def _regenerate_single_fast(
        self,
        qnum: int,
        enhanced_prompt: str,
        storage: ValidQuestionStorage,
        question_type: str
    ) -> str:
        """REGEN NHANH - 1 LẦN THỬ"""
        regen_prompt = self._create_batch_prompt_fast(
            [qnum], enhanced_prompt, storage, question_type
        )
        
        try:
            response = self.client.send_data_to_check(
                prompt=regen_prompt,
                temperature=0.6
            )
            
            validator = QuestionValidator(question_type=question_type)
            questions = validator.parse_questions(response)
            
            return questions.get(qnum, None)
        except:
            return None
    
    def _final_recovery_optimized(
        self,
        failed_nums: List[int],
        enhanced_prompt: str,
        storage: ValidQuestionStorage,
        question_type: str,
        doc: Document
    ) -> List[int]:
        """RECOVERY TỐI ƯU - ÍT RETRY HƠN"""
        final_failed = []
        
        for qnum in failed_nums[:20]:  # Chỉ xử lý 20 câu đầu
            print(f"🔄 Recovery câu {qnum}...")
            
            for attempt in range(1, self.config.max_final_recovery_retry + 1):
                new_qtext = self._regenerate_single_fast(
                    qnum, enhanced_prompt, storage, question_type
                )
                
                if not new_qtext:
                    continue
                
                storage.replace_question(qnum, new_qtext)
                
                try:
                    self._write_question_to_doc_fast(
                        qnum, new_qtext, doc, question_type,
                        image_tracker=self.image_tracker
                    )
                    print(f"   ✅ OK")
                    break
                except:
                    if attempt < self.config.max_final_recovery_retry:
                        continue
                    else:
                        final_failed.append(qnum)
        
        # Phần còn lại bỏ qua
        if len(failed_nums) > 20:
            final_failed.extend(failed_nums[20:])
        
        return final_failed
    
    def _create_batch_prompt_fast(
        self,
        batch_nums: List[int],
        enhanced_prompt: str,
        storage: ValidQuestionStorage,
        question_type: str
    ) -> str:
        """TẠO PROMPT NHANH - NGẮN GỌN"""
        prompt_parts = [
            enhanced_prompt,
            "",
            f"# YÊU CẦU: SINH {len(batch_nums)} CÂU",
            f"Câu số: {', '.join(map(str, batch_nums))}",
            "",
            "## QUY TẮC:",
            "1. ĐÚNG SỐ THỨ TỰ",
            "2. ĐẦY ĐỦ FORMAT",
            "3. KHÔNG LẶP NỘI DUNG",
            "",
            "BẮT ĐẦU:"
        ]
        
        return "\n".join(prompt_parts)