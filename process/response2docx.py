import json
import os
import sys
import threading
import time
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, List, Optional, Any
import zipfile
import subprocess
import re
from tempfile import NamedTemporaryFile
from docx.oxml import parse_xml
import traceback

_FILE_LOCK = threading.RLock()
_OUTPUT_DIR_LOCK = threading.RLock()

def get_app_path():
    """Lấy đường dẫn chứa file .exe hoặc script"""
    if getattr(sys, 'frozen', False):
        return os.path.dirname(sys.executable)
    return os.path.dirname(os.path.abspath(__file__))
def find_pandoc_executable():
    """
    Tìm pandoc.exe theo thứ tự ưu tiên:
    1. Thư mục 'pandoc' cạnh tool (cho bản build)
    2. PATH hệ thống (cho môi trường dev)
    """
    app_path = get_app_path()
    
    # 1. Tìm trong thư mục cục bộ 'pandoc' (ưu tiên cao nhất)
    local_pandoc = os.path.join(app_path, 'pandoc', 'pandoc.exe')
    if os.path.isfile(local_pandoc):
        # print(f"✅ Sử dụng Pandoc cục bộ: {local_pandoc}")
        return local_pandoc
    
    # 2. Fallback: Tìm trong PATH hệ thống (cho dev)
    import shutil
    system_pandoc = shutil.which('pandoc')
    if system_pandoc:
        # print(f"⚠️ Sử dụng Pandoc hệ thống: {system_pandoc}")
        return system_pandoc
    
    # 3. Không tìm thấy
    print("❌ KHÔNG TÌM THẤY PANDOC!")
    return None

def latex_to_omml_via_pandoc(latex_math_dollar):
    """Chuyển đổi LaTeX sang OMML qua Pandoc"""
    pandoc_exe = find_pandoc_executable()
    
    if not pandoc_exe:
        print("❌ Pandoc không khả dụng, bỏ qua equation")
        return None
    
    try:
        # Chuẩn hóa input (loại bỏ ký tự lạ)
        latex_clean = latex_math_dollar.strip()
        
        # Tạo file tạm với encoding UTF-8 BOM để tránh lỗi
        with NamedTemporaryFile(mode='w', suffix=".docx", delete=False, encoding='utf-8') as temp_docx:
            temp_path = temp_docx.name
        
        # Chạy Pandoc với error handling tốt hơn
        result = subprocess.run(
            [pandoc_exe, '--from=latex', '--to=docx', '-o', temp_path],
            input=latex_clean,
            text=True,
            encoding='utf-8',
            capture_output=True,
            timeout=10,  # Timeout 10s để tránh treo
            creationflags=subprocess.CREATE_NO_WINDOW if hasattr(subprocess, 'CREATE_NO_WINDOW') else 0
        )
 
        if result.returncode != 0:
            error_msg = result.stderr.strip() if result.stderr else "Unknown error"
            print(f"⚠️ Pandoc error (code {result.returncode}): {error_msg}")
            
            # Kiểm tra lỗi phổ biến
            if "not found" in error_msg.lower() or "cannot find" in error_msg.lower():
                print("   → Thiếu DLL dependencies. Kiểm tra lại folder pandoc/")
            elif "syntax" in error_msg.lower():
                print(f"   → LaTeX syntax error: {latex_clean[:50]}...")
            
            return None
        
        # Kiểm tra file output có tồn tại không
        if not os.path.exists(temp_path) or os.path.getsize(temp_path) == 0:
            print(f"⚠️ Pandoc không tạo file output hợp lệ")
            return None
           
        # Đọc XML từ DOCX
        with zipfile.ZipFile(temp_path, 'r') as z:
            xml_content = z.read('word/document.xml').decode('utf-8')
        
        # Dọn dẹp file tạm
        try:
            os.remove(temp_path)
        except:
            pass
       
        # Tìm equation XML
        match = re.search(r'(<m:oMath[^>]*>.*?</m:oMath>)', xml_content, re.DOTALL)
        
        if not match:
            print(f"⚠️ Không tìm thấy equation trong output: {latex_clean[:30]}...")
            return None
            
        return match.group(1)
   
    except subprocess.TimeoutExpired:
        print(f"⚠️ Pandoc timeout (>10s)")
        return None
    except Exception as e:
        print(f"❌ Lỗi latex_to_omml: {type(e).__name__}: {e}")
        import traceback
        traceback.print_exc()
        return None



def process_text_with_latex(text, paragraph, bold=False):
    """
    Xử lý text có công thức LaTeX
    VERSION ỔN ĐỊNH - Copy từ test_res.py (KHÔNG có repair_broken_latex)
    """
    if not text:
        return
    
    # Làm sạch HTML tags
    text = text.replace("<br>", "\n").replace("<br/>", "\n") \
               .replace("<Br>", "\n").replace("<Br/>", "\n")
    text = re.sub(r'</?(div|p|u|span|font|i|b)\b[^>]*>', '', text)
    text = text.replace("&nbsp;", "").replace("&lt;", "").replace("&gt;", "")
    
    # Tách text và LaTeX
    pattern = r'(\$[^$]+\$|\\\[.*?\\\])'
    parts = re.split(pattern, text)
    
    for part in parts:
        if not part:
            continue
        
        # Phần LaTeX
        if part.startswith('$') or part.startswith('\\['):
            try:
                latex_expr = clean_latex_math(part)
                insert_equation_into_paragraph(latex_expr, paragraph)
            except Exception as e:
                # Fallback: thêm text thuần
                run = paragraph.add_run(part)
                if bold:
                    run.bold = True
        # Phần text thường
        else:
            cleaned_part = re.sub(r'^\s*/', '', part)
            run = paragraph.add_run(cleaned_part)
            if bold:
                run.bold = True


def insert_equation_into_paragraph(latex_math_dollar, paragraph):
    """Chèn công thức toán học vào paragraph"""
    omml_str = latex_to_omml_via_pandoc(latex_math_dollar)
    
    if not omml_str:
        # Fallback: Thêm text thuần nếu không convert được
        paragraph.add_run(f" [{latex_math_dollar}] ")
        return
    
    # Thêm namespace nếu thiếu
    if 'xmlns:m=' not in omml_str:
        omml_str = re.sub(
            r'<m:oMath',
            r'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"',
            omml_str,
            count=1
        )
    
    try:
        omml_element = parse_xml(omml_str)
        run = paragraph.add_run()
        run._r.append(omml_element)
    except Exception as e:
        print(f"Lỗi chèn equation: {e}")
        paragraph.add_run(f" [{latex_math_dollar}] ")


def clean_latex_math(latex_raw):
    latex_raw = re.sub(r'\\/', '', latex_raw)
    latex_raw = re.sub(r'\\operatorname\s*{\s*([^}]*)\s*}',
                       lambda m: m.group(1).replace(' ', ''), latex_raw)
    latex_raw = re.sub(r'\\root\s*(\d+)\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'\\root\s*{(\d+)}\s*\\of\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'\\root\s*(\d+)\s*\\sqrt\s*{([^}]*)}', r'\\sqrt[\1]{\2}', latex_raw)
    latex_raw = re.sub(r'([a-zA-Z])\s*\\frac\s*{([^}]+)}\s*{([^}]+)}',
                       r'\1^{\\frac{\2}{\3}}', latex_raw)
    latex_raw = re.sub(r'\\sp\s*{([^}]*)}', r'^{\1}', latex_raw)
    latex_raw = re.sub(r'{\\bf\s*([^}]*)}', r'\1', latex_raw)
    latex_raw = re.sub(r'\\\s*log', r'\\log', latex_raw)
    latex_raw = re.sub(r'\\bigskip', '', latex_raw)
    latex_raw = re.sub(r'\\nonumber', '', latex_raw)
    latex_raw = latex_raw.replace(r'\?', '?')
    latex_raw = re.sub(r'\\cdot\s*(?=\w)', r'\\cdot ', latex_raw)
    latex_raw = latex_raw.replace(r'\dotstan', r'\cdot \tan')
    latex_raw = re.sub(r'(?<!\\)(\bln\b|\blog\b|\bsin\b|\bcos\b|\btan\b|\blog_{?\d*}?)',
                       r'\\\1', latex_raw)
    latex_raw = re.sub(r'(\\Leftrightarrow|\\Rightarrow|\\rightarrow)(?=\w)', r'\1 ', latex_raw)
    latex_raw = latex_raw.replace(r'\\n', r'\n')
    
    latex_raw = latex_raw.strip()
    # ✅ KHÁC BIỆT: Version cũ KHÔNG replace \n và \r
    # latex_raw = latex_raw.replace('\n', ' ').replace('\r', '')  # ← XÓA DÒNG NÀY
    
    if not (latex_raw.startswith('$') and latex_raw.endswith('$')):
        latex_raw = f"${latex_raw}$"
    
    return latex_raw

def ensure_output_folder_for_batch(batch_name):
    """Tạo folder riêng cho batch"""
    base_path = get_app_path()
    output_base = os.path.join(base_path, "output")
    batch_folder = os.path.join(output_base, batch_name)
    
    with _OUTPUT_DIR_LOCK:
        os.makedirs(output_base, exist_ok=True)
        os.makedirs(batch_folder, exist_ok=True)
    
    return batch_folder

def save_document_securely(doc, batch_name, file_name):
    """Lưu file DOCX với thread-safety"""
    batch_folder = ensure_output_folder_for_batch(batch_name)
    if not batch_folder:
        return None

    output_path = os.path.join(batch_folder, f"{file_name}.docx")
    
    with _FILE_LOCK:
        max_retries = 3
        for retry_count in range(max_retries):
            try:
                doc.save(output_path)
                if os.path.exists(output_path):
                    file_size = os.path.getsize(output_path)
                    print(f"✅ Đã lưu file: {output_path} ({file_size} bytes)")
                    return output_path
            except Exception as e:
                print(f"⚠️ Lỗi lưu file lần {retry_count + 1}: {e}")
                if retry_count < max_retries - 1:
                    time.sleep(1)
        
        print(f"❌ Không thể lưu file sau {max_retries} lần thử")
        return None
def repair_json_with_ai(broken_json_str: str, client) -> str:
    """
    Gửi JSON lỗi cho AI sửa (ĐÃ CẢI TIẾN)
    Thay vì yêu cầu AI trả về toàn bộ JSON, ta sẽ cố gắng tìm phần hợp lệ
    hoặc yêu cầu AI sửa phần bị lỗi cụ thể.
    """
    print("⚠️ JSON lỗi. Đang cố gắng sửa...")
    
    # --- Cải tiến 1: Cố gắng "xén" JSON hợp lệ ---
    # Đôi khi AI trả về JSON bị thừa văn bản phía sau
    repaired_by_cutting = extract_valid_json(broken_json_str)
    if repaired_by_cutting:
        print("✅ JSON đã được sửa bằng cách xén phần hợp lệ.")
        return repaired_by_cutting

    # --- Cải tiến 2: Gửi yêu cầu sửa lỗi cụ thể hơn cho AI ---
    # Thay vì gửi nguyên đoạn lỗi, mô tả rõ hơn lỗi gì
    error_description = "Đoạn JSON sau bị lỗi cú pháp hoặc bị cắt xén."
    error_description += " Có thể thiếu dấu phẩy, ngoặc đóng/mở không khớp, hoặc bị ngắt giữa chừng."
    error_description += " Vui lòng sửa lỗi cú pháp, GIỮ NGUYÊN TOÀN BỘ NỘI DUNG TIẾNG VIỆT VÀ CÔNG THỨC LATEX,"
    error_description += " và TRẢ VỀ CHỈ CÓ JSON ĐÃ SỬA (không thêm lời dẫn, không thêm markdown)."

    prompt_fix = f"""
{error_description}

JSON lỗi:
{broken_json_str}

JSON đã sửa (chỉ JSON, không thêm gì khác):
"""
    try:
        repaired_text = client.send_data_to_check(prompt_fix)
        # Sau khi AI trả về, thử xén lại lần nữa nếu cần
        final_repaired = extract_valid_json(repaired_text)
        if final_repaired:
            print("✅ JSON đã được sửa bởi AI và xác nhận hợp lệ.")
            return final_repaired
        else:
            print("❌ AI trả về văn bản nhưng vẫn không phải JSON hợp lệ.")
            return repaired_text # Trả về nguyên văn để thử parse sau
    except Exception as e:
        print(f"❌ Gặp lỗi khi yêu cầu AI sửa JSON: {e}")
        return broken_json_str
    
def extract_valid_json(text: str) -> str:
    """
    Cố gắng trích xuất phần JSON hợp lệ từ một chuỗi có thể có văn bản thừa.
    Hỗ trợ cả JSON lồng nhau phức tạp.
    """
    text = text.strip()
    
    # 1. Tìm tất cả các cặp ngoặc nhọn {} hoặc ngoặc vuông []
    stack = []
    start = -1
    max_depth = 0
    max_start = -1
    max_end = -1

    for i, char in enumerate(text):
        if char == '{':
            if not stack:
                start = i
            stack.append(char)
        elif char == '[':
            if not stack:
                start = i
            stack.append(char)
        elif char == '}' and stack and stack[-1] == '{':
            stack.pop()
            if not stack and (i - start) > (max_end - max_start):
                max_start = start
                max_end = i
        elif char == ']' and stack and stack[-1] == '[':
            stack.pop()
            if not stack and (i - start) > (max_end - max_start):
                max_start = start
                max_end = i

    if max_start != -1 and max_end != -1:
        potential_json = text[max_start : max_end + 1]
        # 2. Thử parse phần JSON này
        try:
            # Dùng strict=False để thư giãn một chút với ký tự đặc biệt
            json.loads(potential_json, strict=False)
            print(f"🔍 Đã tìm thấy JSON hợp lệ trong văn bản.")
            return potential_json
        except (json.JSONDecodeError, TypeError):
            pass # Không parse được, thử phương pháp khác

    # Nếu không tìm được JSON hoàn chỉnh, thử phương pháp cũ
    # "Săn" JSON bằng cách tìm { đầu và } cuối
    start_brace = text.find('{')
    end_brace = text.rfind('}')
    if start_brace != -1 and end_brace != -1 and end_brace > start_brace:
        potential_json_old_way = text[start_brace : end_brace + 1]
        try:
            json.loads(potential_json_old_way, strict=False)
            print(f"🔍 Đã tìm thấy JSON hợp lệ theo cách cũ.")
            return potential_json_old_way
        except (json.JSONDecodeError, TypeError):
            pass

    # Nếu tất cả đều không thành công
    return ""
def sanitize_latex_json(text: str) -> str:
    VALID_ESCAPES = {
        '\\\\', '\\"', '\\/', '\\b', '\\f', '\\n', '\\r', '\\t'
    }
    
    def is_already_escaped(s: str, pos: int) -> bool:
        """Kiểm tra backslash tại vị trí pos đã được escape chưa"""
        count = 0
        i = pos - 1
        while i >= 0 and s[i] == '\\':
            count += 1
            i -= 1
        return count % 2 == 1
    
    def fix_string_content(match):
        full = match.group(0)
        content = match.group(1)
        
        if not content:
            return full
        
        result = []
        i = 0
        
        while i < len(content):
            char = content[i]
            
            if char == '\\' and i + 1 < len(content):
                next_char = content[i + 1]
                two_chars = char + next_char
                
                # ✅ Case 1: JSON escape hợp lệ → Giữ nguyên
                if two_chars in VALID_ESCAPES:
                    result.append(two_chars)
                    i += 2
                    continue
                
                # ✅ Case 2: Unicode escape (\uXXXX) → Giữ nguyên
                if next_char == 'u' and i + 5 < len(content):
                    hex_part = content[i+2:i+6]
                    if len(hex_part) == 4 and all(c in '0123456789ABCDEFabcdef' for c in hex_part):
                        result.append(content[i:i+6])
                        i += 6
                        continue
                
                # ✅ Case 3: Backslash ĐÃ escape (\\) → Giữ nguyên
                if next_char == '\\':
                    result.append('\\\\')
                    i += 2
                    continue
                
                # ❌ Case 4: LaTeX command → CẦN escape
                result.append('\\\\')
                i += 1
            else:
                result.append(char)
                i += 1
        
        return '"' + ''.join(result) + '"'
    
    string_pattern = r'"((?:[^"\\]|\\.)*)"'
    sanitized = re.sub(string_pattern, fix_string_content, text)
    return sanitized

def parse_json_safely(json_str: str, client) -> Optional[Dict]:
    """Parse JSON với 3 lớp bảo vệ - ĐÃ CẢI TIẾN"""
    
    # === LỚP 1: Thử parse ngay lập tức ===
    try:
        return json.loads(json_str, strict=False)
    except (json.JSONDecodeError, TypeError):
        pass

    # === LỚP 2: Clean + Sanitize ===
    cleaned = clean_json_string(json_str)
    sanitized = sanitize_latex_json(cleaned)
    
    try:
        return json.loads(sanitized, strict=False)
    except (json.JSONDecodeError, TypeError) as e:
        print(f"⚠️ Lỗi JSON sau sanitize (vị trí {getattr(e, 'pos', 'unknown')}): {e.msg}")
        # Chỉ in context nếu có vị trí lỗi
        if hasattr(e, 'pos'):
            start = max(0, e.pos - 30)
            end = min(len(sanitized), e.pos + 30)
            print(f"Context: ...{sanitized[start:end]}...")

    # === LỚP 3: Gọi AI sửa ===
    print("🤖 Gọi AI sửa JSON...")
    try:
        # Thay vì gửi 'cleaned', ta gửi 'sanitized' để AI có thể sửa dễ hơn
        repaired = repair_json_with_ai(sanitized, client)
        
        # AI có thể trả về nhiều thứ, cố gắng extract JSON hợp lệ
        extracted_json = extract_valid_json(repaired)
        if extracted_json:
            # Nếu extract được, lại phải sanitize lại vì có thể AI chưa escape đúng
            final_sanitized = sanitize_latex_json(extracted_json)
            try:
                return json.loads(final_sanitized, strict=False)
            except (json.JSONDecodeError, TypeError) as e:
                 print(f"❌ AI trả về JSON nhưng vẫn lỗi sau khi sanitize: {e}")
                 return None
        else:
            # Nếu không extract được, thử parse nguyên văn
            try:
                return json.loads(repaired, strict=False)
            except (json.JSONDecodeError, TypeError) as e:
                 print(f"❌ Không thể parse cả JSON được AI sửa: {e}")
                 return None

    except Exception as e:
        print(f"❌ LỖI nghiêm trọng khi gọi AI sửa JSON: {e}")
        traceback.print_exc()
        return None


def clean_json_string(text: str) -> str:
    """Xóa markdown wrapper và lấy phần JSON thuần túy"""
    if not text:
        return ""
    
    text = text.strip()
    
    # Xóa markdown code block
    pattern = r"```(?:json)?(.*?)```"
    match = re.search(pattern, text, re.DOTALL | re.IGNORECASE)
    if match:
        text = match.group(1).strip()
    
    # "Săn" JSON bằng cách tìm { đầu và } cuối
    start = text.find('{')
    end = text.rfind('}')
    
    if start != -1 and end != -1 and end > start:
        return text[start:end + 1]
    
    return text
def generate_or_get_image(hinh_anh_data: Dict) -> tuple:
    """
    Xử lý gọi hàm sinh ảnh.
    Returns: (image_bytes, placeholder_text) - image_bytes là 1 object duy nhất
    """
    mo_ta = hinh_anh_data.get("mo_ta", hinh_anh_data.get("description", ""))
    mo_ta = str(mo_ta).strip()
    loai = hinh_anh_data.get("loai", "tu_mo_ta")
    
    if loai == "tu_mo_ta" and mo_ta:
        try:
            from process.text2Image import generate_image_from_text
            # Hàm này trả về 1 bytes object (hoặc None)
            image_bytes = generate_image_from_text(mo_ta)
            if image_bytes:
                return image_bytes, None
            else:
                # Nếu API trả về None (do lỗi mạng hoặc quota)
                return None, f"⚠️ [Lỗi sinh ảnh] Server không trả về ảnh cho mô tả: {mo_ta}"
        except Exception as e:
            print(f"❌ Lỗi sinh ảnh: {e}")
            return None, f"⚠️ [Lỗi Code] {str(e)}"
    
    placeholder = f"🖼️ [Cần chèn hình: {mo_ta}]"
    return None, placeholder

def insert_image_or_placeholder(doc: Document, hinh_anh_data: Dict):
    """Chèn ảnh hoặc placeholder vào document"""
    image_bytes, placeholder = generate_or_get_image(hinh_anh_data)
    
    if image_bytes:
        try:
            image_stream = BytesIO(image_bytes)
            doc.add_picture(image_stream, width=Inches(4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"❌ Lỗi chèn ảnh: {e}")
            p = doc.add_paragraph()
            run = p.add_run(f"⚠️ [Lỗi chèn ảnh: {str(e)}]")
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.italic = True
    
    elif placeholder:
        p = doc.add_paragraph()
        run = p.add_run(placeholder)
        run.font.color.rgb = RGBColor(200, 0, 0)
        run.italic = True
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc

class PromptBuilder:
    """
    Builder tạo prompt động - Đã sửa lỗi xung đột f-string và LaTeX
    """
    
    @staticmethod
    def build_json_structure_hint(question_type: str) -> str:
        # Hàm này trả về string thường (không phải f-string) nên giữ nguyên 1 dấu {}
        if question_type == "trac_nghiem_4_dap_an":
            return """
{
  "loai_de": "trac_nghiem_4_dap_an",
  "tong_so_cau": 80,
  "cau_hoi": [
    {
      "stt": 1,
      "muc_do": "nhan_biet",
      "phan": "Phần I",
      "noi_dung": "Nội dung câu hỏi...",
      "hinh_anh": { "co_hinh": true, "loai": "tu_mo_ta", "mo_ta": "Mô tả..." },
      "dap_an": [
        {"ky_hieu": "A", "noi_dung": "Đáp án A"},
        {"ky_hieu": "B", "noi_dung": "Đáp án B"},
        {"ky_hieu": "C", "noi_dung": "Đáp án C"},
        {"ky_hieu": "D", "noi_dung": "Đáp án D"}
      ],
      "dap_an_dung": 2,
      "giai_thich": "Giải thích..."
    }
  ]
}
"""
        elif question_type == "dung_sai":
            return """
{
  "loai_de": "dung_sai",
  "tong_so_cau": 40,
  "cau_hoi": [
    {
      "stt": 1,
      "muc_do": "thong_hieu",
      "phan": "Phần I",
      "doan_thong_tin": "...",
      "hinh_anh": { "co_hinh": true, "loai": "tu_mo_ta", "mo_ta": "Mô tả..." },
      "cac_y": [
        {"ky_hieu": "a", "noi_dung": "...", "dung": false},
        {"ky_hieu": "b", "noi_dung": "...", "dung": true},
        {"ky_hieu": "c", "noi_dung": "...", "dung": false},
        {"ky_hieu": "d", "noi_dung": "...", "dung": true}
      ],
      "dap_an_dung_sai": "0101",
      "giai_thich": [
        {"y": "a", "noi_dung_y": "...", "ket_luan": "SAI", "giai_thich": "..."},
        {"y": "b", "noi_dung_y": "...", "ket_luan": "ĐÚNG", "giai_thich": "..."}
      ]
    }
  ]
}
"""
        elif question_type == "tra_loi_ngan":
            return """
{
  "loai_de": "tra_loi_ngan",
  "tong_so_cau": 30,
  "cau_hoi": [
    {
      "stt": 1,
      "muc_do": "van_dung",
      "phan": "Phần I",
      "noi_dung": "...",
      "hinh_anh": { "co_hinh": true, "loai": "tu_mo_ta", "mo_ta": "..." },
      "dap_an": "[[kết quả]]",
      "giai_thich": "..."
    }
  ]
}
"""
        return "{}"

    @staticmethod
    def wrap_user_prompt(user_prompt: str, question_type: str) -> str:
        json_hint = PromptBuilder.build_json_structure_hint(question_type)
        
        # SỬA LỖI TẠI ĐÂY:
        # Trong f-string (f"""), dấu ngoặc nhọn của LaTeX phải nhân đôi thành {{ }}
        # Ví dụ: \frac{1}{2} phải viết là \\frac{{1}}{{2}} để Python không hiểu nhầm là biến
        
        return f"""{user_prompt}

----------------
### YÊU CẦU NGHIÊM NGẶT VỀ DỮ LIỆU (BẮT BUỘC TUÂN THỦ 100%):

1. **FORMAT ĐẦU RA (QUAN TRỌNG NHẤT)**: 
   - CHỈ TRẢ VỀ DUY NHẤT MỘT CHUỖI JSON thuần túy.
   - TUYỆT ĐỐI KHÔNG có lời mở đầu hay kết thúc (như "Here is result...").
   - **QUAN TRỌNG:** Phải sử dụng dấu ngoặc kép (") cho key và value. KHÔNG dùng dấu ngoặc đơn (').
2. **HÌNH ẢNH:**
   - Nếu câu hỏi có hình, BẮT BUỘC điền mô tả chi tiết vào trường "hinh_anh".
   - Ví dụ: "Tam giác ABC vuông tại A..." hoặc "Sơ đồ mạch điện gồm..."

### MẪU JSON MONG MUỐN:
{json_hint}
"""
class DynamicDocxRenderer:
    """
    Renderer tự động thích ứng với cấu trúc JSON
    """
    
    def __init__(self, doc: Document):
        self.doc = doc
    
    def render_title(self, data: Dict):
        """Render tiêu đề tự động"""
        loai_de = data.get("loai_de", "").upper()
        title = self.doc.add_heading(f'ĐỀ {loai_de}', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def auto_group_questions(self, data: Dict) -> Dict[str, List]:
        """
        Tự động nhóm câu hỏi và CHUẨN HÓA key muc_do từ Tiếng Việt sang code.
        Giúp người dùng thoải mái viết prompt "Vận dụng", "Nhận biết"... mà không bị lỗi file trắng.
        """
        grouped = {}
        for cau in data.get("cau_hoi", []):
            # 1. Lấy dữ liệu thô từ AI (ví dụ: "Vận dụng", "Nhận biết", "Thông hiểu")
            # Chuyển về chữ thường để dễ so sánh
            raw_muc_do = str(cau.get("muc_do", "unknown")).lower().strip()
            
            # 2. Logic "Phiên dịch" thông minh (Mapping)
            # Ưu tiên check "cao" trước để phân biệt "Vận dụng" và "Vận dụng cao"
            if "cao" in raw_muc_do:
                muc_do_chuan = "van_dung_cao"
            elif "dụng" in raw_muc_do or "dung" in raw_muc_do:
                muc_do_chuan = "van_dung"
            elif "thông" in raw_muc_do or "thong" in raw_muc_do:
                muc_do_chuan = "thong_hieu"
            elif "nhận" in raw_muc_do or "nhan" in raw_muc_do:
                muc_do_chuan = "nhan_biet"
            else:
                # Trường hợp AI ghi nội dung lạ, mặc định đưa vào Vận dụng 
                # để đảm bảo câu hỏi vẫn hiện ra trong file (tránh lỗi trang trắng)
                muc_do_chuan = "van_dung" 
            
            # 3. Gom nhóm theo key chuẩn
            if muc_do_chuan not in grouped:
                grouped[muc_do_chuan] = []
            grouped[muc_do_chuan].append(cau)
        
        # Sắp xếp theo STT trong mỗi nhóm
        for key in grouped:
            grouped[key].sort(key=lambda x: x.get("stt", 0))
        
        return grouped
    
    def get_section_title(self, muc_do: str) -> str:
        """
        Tạo tiêu đề section dựa trên mức độ
        CÓ THỂ mở rộng bằng config file
        """
        mapping = {
            "nhan_biet": "I. CÂU HỎI NHẬN BIẾT",
            "thong_hieu": "II. CÂU HỎI THÔNG HIỂU",
            "van_dung": "III. CÂU HỎI VẬN DỤNG",
            "van_dung_cao": "IV. CÂU HỎI VẬN DỤNG CAO"
        }
        return mapping.get(muc_do, muc_do.upper())
    
    def render_question_trac_nghiem(self, cau: Dict):
        """Render câu hỏi trắc nghiệm 4 đáp án"""
        # Câu hỏi
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}. ").bold = True
        process_text_with_latex(cau['noi_dung'], p) 
        
        # Hình ảnh
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
        
        # Đáp án - THÊM XỬ LÝ LATEX
        for dap_an in cau.get("dap_an", []):
            p_da = self.doc.add_paragraph()
            run_ky_hieu = p_da.add_run(f"{dap_an['ky_hieu']}. ")
            process_text_with_latex(dap_an['noi_dung'], p_da) 
        
        # Lời giải
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
        
        if "dap_an_dung" in cau:
            p_dung = self.doc.add_paragraph()
            p_dung.add_run(f"{cau['dap_an_dung']}").bold = True
            self.doc.add_paragraph("####")
        
        # Giải thích - THÊM XỬ LÝ LATEX
        giai_thich = cau.get("giai_thich", "")
        for line in giai_thich.split("\n"):
            if line.strip():
                p_gt = self.doc.add_paragraph()
                process_text_with_latex(line.strip(), p_gt)  
        
        # Kết luận - THÊM XỬ LÝ LATEX
        if "dap_an_dung" in cau:
            dap_an_num = cau['dap_an_dung']
            noi_dung_dap_an = cau['dap_an'][dap_an_num-1]['noi_dung']
            p_ket_luan = self.doc.add_paragraph()
            run = p_ket_luan.add_run("Vậy đáp án đúng là: ")
            run.bold = True
            process_text_with_latex(noi_dung_dap_an, p_ket_luan, bold=True) 
    
    def render_question_dung_sai(self, cau: Dict):
        """Render câu hỏi đúng/sai"""
        # Số câu
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}.").bold = True
        
        # Đoạn thông tin - THÊM XỬ LÝ LATEX
        if cau.get("doan_thong_tin"):
            p_doan = self.doc.add_paragraph()
            process_text_with_latex(cau.get("doan_thong_tin", ""), p_doan)  
        
        # Hình ảnh
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
        
        # Các ý a, b, c, d - THÊM XỬ LÝ LATEX
        for y in cau.get("cac_y", []):
            p_y = self.doc.add_paragraph()
            p_y.add_run(f"{y['ky_hieu']}) ")
            process_text_with_latex(y['noi_dung'], p_y)  
        
        # Lời giải
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
        
        p_da = self.doc.add_paragraph()
        p_da.add_run(cau.get("dap_an_dung_sai", "")).bold = True
        self.doc.add_paragraph("####")
        
        # Giải thích từng ý - THÊM XỬ LÝ LATEX
        for gt in cau.get("giai_thich", []):
            p_gt = self.doc.add_paragraph()
            p_gt.add_run('+) "') 
            process_text_with_latex(gt.get('noi_dung_y', ''), p_gt)
            ket_luan = gt.get('ket_luan', 'SAI')
            run_kl = p_gt.add_run(f'" - {ket_luan}. ')
            run_kl.bold = True
            
            if gt.get('giai_thich'):
                # p_gt_detail = self.doc.add_paragraph()
                process_text_with_latex(gt.get('giai_thich', ''), p_gt)  
    
    def render_question_tra_loi_ngan(self, cau: Dict):
        """Render câu hỏi trả lời ngắn"""
        # Câu hỏi
        p = self.doc.add_paragraph()
        p.add_run(f"Câu {cau['stt']}. ").bold = True
        p_noi_dung = self.doc.add_paragraph()
        process_text_with_latex(cau['noi_dung'], p_noi_dung)  
        
        # Hình ảnh (nếu có)
        hinh_anh = cau.get("hinh_anh", {})
        if hinh_anh.get("co_hinh"):
            insert_image_or_placeholder(self.doc, hinh_anh)
        
        # Đáp án - THÊM XỬ LÝ LATEX
        p_da = self.doc.add_paragraph()
        run_label = p_da.add_run("Đáp án: ")
        run_label.bold = True
        
        raw_ans = str(cau.get('dap_an', '')).strip()
        if raw_ans.startswith("[[") and raw_ans.endswith("]]"):
            final_ans = raw_ans
        else:
            final_ans = f"[[{raw_ans}]]"
        
        # XỬ LÝ LATEX TRONG ĐÁP ÁN
        process_text_with_latex(final_ans, p_da, bold=True)  
        
        # Lời giải header
        p_lg = self.doc.add_paragraph()
        p_lg.add_run("Lời giải").bold = True
        self.doc.add_paragraph("####")
        
        # Giải thích chi tiết - ĐÃ CÓ XỬ LÝ LATEX
        giai_thich = cau.get("giai_thich", "")
        lines = giai_thich.replace('\\n', '\n').split('\n')
        
        for line in lines:
            text = line.strip()
            if not text or text == "####":
                continue
            
            is_bold = False
            if text.startswith("**") and text.endswith("**"):
                text = text[2:-2]
                is_bold = True
            
            check_text = text.replace('*', '').strip().lower()
            if check_text.startswith("vậy"):
                is_bold = True
                text = text.replace('**', '')

            p_gt = self.doc.add_paragraph()
            process_text_with_latex(text, p_gt, bold=is_bold)  
    
    def render_all(self, data: Dict):
        """
        Main render function - Có hỗ trợ chia PHẦN (PART) bên trong Mức độ
        """
        self.render_title(data)
        
        # 1. Auto-group theo mức độ (Nhận biết, Thông hiểu...)
        grouped = self.auto_group_questions(data)
        
        # 2. Detect loại đề
        loai_de = data.get("loai_de", "")
        
        # 3. Render từng nhóm MỨC ĐỘ
        # Thứ tự ưu tiên render
        order_muc_do = ["nhan_biet", "thong_hieu", "van_dung", "van_dung_cao"]
        
        for muc_do in order_muc_do:
            if muc_do not in grouped:
                continue
            
            # Lấy danh sách câu hỏi trong mức độ này
            questions = grouped[muc_do]
            if not questions:
                continue
            section_title = self.get_section_title(muc_do)
            self.doc.add_heading(section_title, level=2)
            current_phan = None

            for cau in questions:
                # Lấy tên phần của câu hiện tại
                phan_cua_cau = str(cau.get("phan", "")).strip()
                
                # Nếu câu này thuộc một phần mới -> In Header Phần
                if phan_cua_cau and phan_cua_cau != current_phan:
                    # In ra header cấp 3 (VD: Phần 1: Đội ngũ...)
                    # Dùng màu hoặc in đậm để phân biệt
                    p_phan = self.doc.add_heading(phan_cua_cau.upper(), level=3)
                    p_phan.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    current_phan = phan_cua_cau
                
                # Render nội dung câu hỏi như bình thường
                if loai_de == "dung_sai":
                    self.render_question_dung_sai(cau)
                elif loai_de == "tra_loi_ngan":
                    self.render_question_tra_loi_ngan(cau)
                else:
                    self.render_question_trac_nghiem(cau)

def response2docx_flexible(
    file_path: str,
    prompt: str,
    file_name: str,
    project_id: str,
    creds: str,
    model_name: str,
    question_type: str = "trac_nghiem_4_dap_an",
    batch_name: Optional[str] = None
) -> Optional[str]:
    """
    Phiên bản cải tiến với cơ chế fail-safe triệt để.
    Luôn luôn trả về 1 file .docx, dù AI có lỗi hay không.
    """
    try:
        from api.callAPI import VertexClient
        from docx import Document # Import ở đây để đảm bảo nếu lỗi ở mức import thì vẫn bắt được
        
        client = VertexClient(project_id, creds, model_name)
        
        if not batch_name:
            batch_name = file_name.replace("_TN", "").replace("_DS", "").replace("_TLN", "")
        
        # 1. Wrap prompt với JSON structure hint
        final_prompt = PromptBuilder.wrap_user_prompt(prompt, question_type)
        
        # 2. Gửi request AI
        print("📤 Đang gửi request tới AI...")
        ai_response = client.send_data_to_AI(final_prompt, file_path)
        
        # 3. Parse JSON với cơ chế an toàn
        print("🔄 Đang parse JSON...")
        data = parse_json_safely(ai_response, client)
        
        if not data:
            print("❌ Không thể parse JSON từ AI. Sử dụng chế độ FAIL-SAFE.")
            # --- BƯỚC FAIL-SAFE TRIỆT ĐỂ ---
            # Tạo một file .docx tối thiểu với nội dung phản hồi thô từ AI
            doc = Document()
            doc.add_heading(f'ĐỀ {question_type.upper()}', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_heading('PHẢN HỒI TỪ AI (RAW)', level=2)
            # Thêm phản hồi thô để người dùng biết AI đã trả về gì
            doc.add_paragraph(ai_response)
            
            doc.add_heading('LỖI', level=3)
            doc.add_paragraph('Dữ liệu từ AI không ở định dạng JSON hợp lệ và không thể xử lý.')
            doc.add_paragraph('Vui lòng kiểm tra lại prompt hoặc nội dung file đầu vào.')
            
            print("📝 Đang lưu file FAIL-SAFE...")
            output_path = save_document_securely(doc, batch_name, f"{file_name}_loi_parse")
            if output_path:
                print(f"✅ Đã lưu file FAIL-SAFE: {output_path}")
            else:
                print("❌ Không thể lưu file FAIL-SAFE.")
            return output_path # Trả về đường dẫn file fail-safe

        print(f"✅ Parse thành công: {data.get('tong_so_cau', 0)} câu hỏi")
        
        # 4. Render DOCX động (cũng có thể gây lỗi)
        print("📝 Đang tạo DOCX...")
        doc = Document()
        renderer = DynamicDocxRenderer(doc)
        
        try:
            renderer.render_all(data)
            print("✅ Render DOCX thành công")
        except Exception as e_render:
            print(f"❌ Lỗi khi render DOCX: {e_render}")
            print("📝 Đang chuyển sang chế độ FAIL-SAFE (dữ liệu thô)...")
            traceback.print_exc() # Ghi log lỗi chi tiết
            
            # --- BƯỚC FAIL-SAFE CHO RENDER ---
            # Tạo lại document mới từ đầu, chỉ ghi dữ liệu thô
            doc = Document()
            doc.add_heading(f'ĐỀ {question_type.upper()}', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_heading('DỮ LIỆU TỪ AI (RAW - JSON)', level=2)
            # Chuyển dữ liệu sang string và thêm vào doc
            raw_data_str = json.dumps(data, ensure_ascii=False, indent=2)
            doc.add_paragraph(raw_data_str)
            
            doc.add_heading('LỖI KHI XỬ LÝ', level=3)
            doc.add_paragraph(f'Lỗi render: {e_render}')
            doc.add_paragraph('Dữ liệu thô đã được lưu. Vui lòng kiểm tra lại cấu trúc JSON.')
            
            print("📝 Đang lưu file FAIL-SAFE (render lỗi)...")
            output_path = save_document_securely(doc, batch_name, f"{file_name}_loi_render")
            if output_path:
                print(f"✅ Đã lưu file FAIL-SAFE (render lỗi): {output_path}")
            else:
                print("❌ Không thể lưu file FAIL-SAFE (render lỗi).")
            return output_path # Trả về đường dẫn file fail-safe

        # 5. Lưu file (có thể lỗi do quyền truy cập, ổ đĩa đầy, v.v.)
        print("💾 Đang lưu file...")
        output_path = save_document_securely(doc, batch_name, file_name)
        
        if output_path:
            print(f"✅ Hoàn thành: {output_path}")
        else:
            print("❌ Không thể lưu file (lỗi từ hàm save_document_securely).")
            # --- BƯỚC FAIL-SAFE CHO LƯU FILE ---
            # Không thể lưu theo tên gốc, thử lưu với tên lỗi
            print("📝 Đang chuyển sang chế độ FAIL-SAFE (lỗi lưu file)...")
            fallback_doc_path = os.path.join(ensure_output_folder_for_batch(batch_name), f"{file_name}_loi_luu.docx")
            try:
                doc.save(fallback_doc_path)
                print(f"✅ Đã lưu file FAIL-SAFE (lỗi lưu file): {fallback_doc_path}")
                return fallback_doc_path
            except Exception as e_save:
                print(f"❌ FAIL-SAFE cũng thất bại khi lưu: {e_save}")
                return None # Cuối cùng vẫn thất bại
            
        return output_path

    except Exception as e_main:
        # BƯỚC FAIL-SAFE CUỐI CÙNG CHO TOÀN BỘ HÀM
        print(f"❌ LỖI NGHIÊM TRỌNG TRONG TOÀN BỘ HÀM: {e_main}")
        traceback.print_exc()
        
        # Tạo một file .docx trống tối thiểu với thông báo lỗi
        try:
            doc = Document()
            doc.add_heading('LỖI HỆ THỐNG', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph(f'Lỗi nghiêm trọng: {e_main}')
            doc.add_paragraph('Hệ thống không thể xử lý yêu cầu.')
            
            if not batch_name:
                batch_name = file_name.replace("_TN", "").replace("_DS", "").replace("_TLN", "")
            fallback_path = os.path.join(ensure_output_folder_for_batch(batch_name), f"{file_name}_loi_he_thong.docx")
            
            doc.save(fallback_path)
            print(f"✅ Đã tạo file FAIL-SAFE cuối cùng: {fallback_path}")
            return fallback_path
        except Exception as e_final:
            print(f"❌ Không thể tạo file FAIL-SAFE cuối cùng: {e_final}")
            return None

def response2docx_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho trắc nghiệm 4 đáp án (legacy)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="trac_nghiem_4_dap_an",
        batch_name=batch_name
    )

def response2docx_dung_sai_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho đúng/sai (legacy)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="dung_sai",
        batch_name=batch_name
    )
    
def response2docx_tra_loi_ngan_json(file_path, prompt, file_name, project_id, creds, model_name, batch_name=None):
    """Wrapper cho trả lời ngắn (legacy compatibility)"""
    return response2docx_flexible(
        file_path, prompt, file_name, project_id, creds, model_name,
        question_type="tra_loi_ngan",
        batch_name=batch_name
    )

class ConfigManager:
    """
    Quản lý cấu hình qua file JSON
    Cho phép thay đổi TOÀN BỘ behavior mà không sửa code
    """
    
    DEFAULT_CONFIG = {
        "section_mapping": {
            "nhan_biet": "I. CÂU HỎI NHẬN BIẾT",
            "thong_hieu": "II. CÂU HỎI THÔNG HIỂU",
            "van_dung": "III. CÂU HỎI VẬN DỤNG",
            "van_dung_cao": "IV. CÂU HỎI VẬN DỤNG CAO" 
        },
        "section_order": ["nhan_biet", "thong_hieu", "van_dung", "van_dung_cao"],
        "auto_fix": True,
        "image_width_inches": 4,
        "retry_json_parse": 2
    }
    
    @classmethod
    def load_config(cls, config_path: str = "config.json") -> Dict:
        """Load config từ file hoặc dùng default"""
        if os.path.exists(config_path):
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        return cls.DEFAULT_CONFIG
    
    @classmethod
    def save_config(cls, config: Dict, config_path: str = "config.json"):
        """Lưu config để tái sử dụng"""
        with open(config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)